"""Parses an uploaded RFP/questionnaire into raw text.

Handles two ways the questionnaire can arrive in `adk web`: as a file
attached to the chat turn (inline data on the user's message) or as text
pasted directly into the chat. Format for attachments is decided by MIME
type, not content sniffing: RFPs get uploaded with whatever type the browser
reports, and guessing from bytes buys nothing for a demo corpus of one file
type per run.
"""

from __future__ import annotations

import csv
import io

from google.adk.tools import ToolContext

_MIME_TO_EXTENSION = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
    "application/pdf": ".pdf",
    "text/csv": ".csv",
    "text/plain": ".txt",
    "text/markdown": ".md",
}


def _parse_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_xlsx(data: bytes) -> str:
    import openpyxl

    workbook = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    parts = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _parse_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    rows = csv.reader(io.StringIO(text))
    return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)


def _parse_text(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


_PARSERS = {
    ".docx": _parse_docx,
    ".xlsx": _parse_xlsx,
    ".pdf": _parse_pdf,
    ".csv": _parse_csv,
    ".txt": _parse_text,
    ".md": _parse_text,
}


async def parse_document(tool_context: ToolContext) -> dict:
    """Reads the questionnaire from the current chat turn and returns its text.

    Checks, in order: an inline file attached to this message, a saved
    session artifact (if the client uploads that way instead), then falls
    back to any pasted text already present in the message.

    Sets skip_summarization on every path so this is always the turn's only
    model call: the LLM is forced to call this tool (see agent.py), and
    without skip_summarization that would also force a second, pointless
    tool-call attempt on the follow-up summarization turn.

    Args:
        tool_context: injected by ADK, gives access to the current turn and
            session artifacts.

    Returns:
        A dict with "source" and "characters_parsed", or "error" if nothing
        usable was found. The parsed text itself goes to state, not the
        returned dict, so it never gets echoed into the chat pane.
    """
    tool_context.actions.skip_summarization = True

    user_content = tool_context.user_content
    if user_content is not None and user_content.parts:
        for part in user_content.parts:
            if part.inline_data is not None and part.inline_data.data:
                extension = _MIME_TO_EXTENSION.get(part.inline_data.mime_type or "")
                if extension is None:
                    return {
                        "error": (
                            f"Unsupported attachment type "
                            f"{part.inline_data.mime_type!r}. Supported: "
                            f"{', '.join(sorted(set(_MIME_TO_EXTENSION.values())))}."
                        )
                    }
                raw_text = _PARSERS[extension](part.inline_data.data)
                if raw_text.strip():
                    tool_context.state["raw_text"] = raw_text
                    tool_context.state["source_filename"] = f"attachment{extension}"
                    return {"source": "attachment", "characters_parsed": len(raw_text)}

        pasted_text = "\n".join(p.text for p in user_content.parts if p.text)
        if pasted_text.strip():
            tool_context.state["raw_text"] = pasted_text
            tool_context.state["source_filename"] = "pasted_text"
            return {"source": "pasted_text", "characters_parsed": len(pasted_text)}

    artifacts = await tool_context.list_artifacts()
    if artifacts:
        filename = artifacts[0]
        extension = (
            "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        )
        parser = _PARSERS.get(extension)
        if parser is None:
            return {
                "error": (
                    f"Unsupported file type {extension!r} for {filename!r}. "
                    f"Supported: {', '.join(sorted(set(_MIME_TO_EXTENSION.values())))}."
                )
            }
        part = await tool_context.load_artifact(filename)
        if part is not None and part.inline_data is not None and part.inline_data.data:
            raw_text = parser(part.inline_data.data)
            if raw_text.strip():
                tool_context.state["raw_text"] = raw_text
                tool_context.state["source_filename"] = filename
                return {"source": filename, "characters_parsed": len(raw_text)}

    return {
        "error": (
            "No questionnaire found. Attach a docx/xlsx/pdf/csv/txt/md file, "
            "or paste the questionnaire text directly in chat."
        )
    }
